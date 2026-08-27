from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


VERSION = "0.2"
DATE = "27.08.2026"


def set_cell_text(cell, value: str) -> None:
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)


def insert_before(anchor, paragraph) -> None:
    anchor._p.addprevious(paragraph._p)


def revise(source: Path, output: Path) -> None:
    document = Document(source)

    # Metadata: the version is fixed for development, but the approval table
    # remains unsigned until the user explicitly fills it in.
    set_cell_text(
        document.tables[0].cell(1, 1),
        "Зафиксированная рабочая редакция для разработки; лист утверждения не заполнен",
    )
    set_cell_text(document.tables[0].cell(2, 1), f"{VERSION} от {DATE}")
    set_cell_text(document.tables[-1].cell(5, 1), VERSION)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            for run in paragraph.runs:
                if "Черновик v0.1" in run.text:
                    run.text = run.text.replace(
                        "Черновик v0.1", f"Рабочая редакция v{VERSION}"
                    )

    end_paragraph = document.paragraphs[-1]
    if "Конец документа" not in end_paragraph.text:
        raise RuntimeError("Не найден завершающий абзац контракта")

    heading = document.add_paragraph(
        "Приложение B. Обязательное применение контракта и регистрация ошибок",
        style="Heading 1",
    )
    insert_before(end_paragraph, heading)

    requirements = (
        "Перед любым изменением кода, сборкой, переносом файлов или тестовым прогоном координатор обязан прочитать текущую версию контракта и указать затрагиваемые требования.",
        "Каждая задача, передаваемая Codex, Spark или агенту, должна содержать путь к текущему контракту, его версию и SHA-256, применимые пункты и проверяемые условия PASS.",
        "Версии контракта неизменяемы. Новое требование или исправление выпускается новой версией; файл CURRENT.md указывает единственную действующую рабочую редакцию.",
        "Каждая ошибка, о которой сообщает пользователь, регистрируется отдельной записью: наблюдаемое поведение, ожидаемое поведение, доказательство, затронутые пункты контракта и обязательный регрессионный тест.",
        "До изменения кода ошибка должна быть отражена в реестре и, если она меняет или уточняет требование, в новой версии контракта. Устное исключение не заменяет зафиксированное требование.",
        "Новая реализация не может отменять ранее принятые функции. При объединении веток и пакетов проверяются все предыдущие условия приёмки, а не только текущая ошибка.",
        "Результат задачи не принимается без доказательства выполнения относящихся к ней пунктов контракта и без сохранённого протокола проверки.",
    )
    for requirement in requirements:
        paragraph = document.add_paragraph(f"— {requirement}", style="Contract List")
        insert_before(end_paragraph, paragraph)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    revise(args.source, args.output)


if __name__ == "__main__":
    main()
